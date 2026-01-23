import streamlit as st
import requests
import pandas as pd
from datetime import datetime
import asyncio
import os
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches
from io import BytesIO

st.set_page_config(page_title="網站監控工具", page_icon="🌐", layout="wide")

st.title("🌐 網站狀態監控 (含截圖功能)")
st.write("點擊下方按鈕以檢查 `urlList.txt` 中的網站狀態並拍攝截圖。")

# Ensure playwright is installed on first run
@st.cache_resource
def install_playwright():
    try:
        # Check if chromium is already installed
        st.info("正在檢查並安裝 Playwright 瀏覽器...")
        os.system("playwright install chromium")
    except Exception as e:
        st.error(f"Playwright 安裝失敗: {e}")

if 'playwright_installed' not in st.session_state:
    install_playwright()
    st.session_state['playwright_installed'] = True

def load_urls(filename='urlList.txt'):
    urls = []
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        current_name = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('http://') or line.startswith('https://'):
                if current_name:
                    urls.append({'Name': current_name, 'URL': line})
                    current_name = None
            else:
                current_name = line
        return urls
    except FileNotFoundError:
        st.error(f"找不到檔案: {filename}")
        return []

async def capture_screenshot(url, filename):
    try:
        async with async_playwright() as p:
            # Add args for better compatibility in containerized environments
            browser = await p.chromium.launch(
                headless=True,
                args=[
                    "--no-sandbox",
                    "--disable-setuid-sandbox",
                    "--disable-dev-shm-usage",
                    "--disable-gpu"
                ]
            )
            context = await browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            )
            page = await context.new_page()
            await page.set_viewport_size({"width": 1280, "height": 720})
            
            # Increase timeout and use a more reliable wait strategy
            await page.goto(url, timeout=45000, wait_until="load")
            # Wait a bit more for dynamic content
            await asyncio.sleep(2)
            
            await page.screenshot(path=filename, full_page=False)
            await browser.close()
            return True
    except Exception as e:
        print(f"截圖失敗 {url}: {e}")
        return False

def check_website(url):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=20, verify=False)
        return {
            'Status': response.status_code,
            'Time (ms)': round(response.elapsed.total_seconds() * 1000, 2),
            'Result': '✅ 正常' if response.status_code == 200 else f'❌ 異常 ({response.status_code})'
        }
    except Exception as e:
        return {
            'Status': 'Error',
            'Time (ms)': 0,
            'Result': f'⚠️ 無法連線'
        }

def create_docx(df):
    doc = Document()
    doc.add_heading('網站狀態監控報告', 0)
    
    doc.add_paragraph(f'報告產生時間: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add summary table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '名稱'
    hdr_cells[1].text = '網址'
    hdr_cells[2].text = '狀態'
    hdr_cells[3].text = '回應時間 (ms)'
    
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['名稱'])
        row_cells[1].text = str(row['網址'])
        row_cells[2].text = str(row['狀態'])
        row_cells[3].text = str(row['回應時間 (ms)'])
        
    doc.add_page_break()
    doc.add_heading('📸 網頁截圖詳情', level=1)
    
    for idx, row in df.iterrows():
        doc.add_heading(row['名稱'], level=2)
        doc.add_paragraph(f"網址: {row['網址']}")
        doc.add_paragraph(f"狀態: {row['狀態']}")
        
        if row['截圖檔案'] and os.path.exists(row['截圖檔案']):
            try:
                doc.add_picture(row['截圖檔案'], width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"無法嵌入圖片: {e}")
        else:
            doc.add_paragraph("⚠️ 無法取得截圖")
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

if st.button('開始檢查 🚀'):
    urls = load_urls()
    
    if not urls:
        st.warning("沒有找到網址或 urlList.txt 為空。")
    else:
        results = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Create a folder for screenshots
        if not os.path.exists('screenshots'):
            os.makedirs('screenshots')
            
        for i, item in enumerate(urls):
            status_text.text(f"正在檢查: {item['Name']} ({item['URL']})...")
            
            # Check status
            check_result = check_website(item['URL'])
            current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Capture screenshot
            screenshot_file = f"screenshots/{i}.png"
            shot_success = asyncio.run(capture_screenshot(item['URL'], screenshot_file))
            
            result_entry = {
                '名稱': item['Name'],
                '網址': item['URL'],
                '測試時間': current_time,
                '狀態': check_result['Result'],
                '回應時間 (ms)': check_result['Time (ms)'],
                '代碼': check_result['Status'],
                '截圖檔案': screenshot_file if shot_success else None
            }
            results.append(result_entry)
            progress_bar.progress((i + 1) / len(urls))
            
        status_text.text("檢查完成！")
        
        df = pd.DataFrame(results)
        st.dataframe(df.drop(columns=['截圖檔案']))

        st.subheader("📸 網頁截圖預覽")
        cols = st.columns(3)
        for idx, row in df.iterrows():
            with cols[idx % 3]:
                if row['截圖檔案'] and os.path.exists(row['截圖檔案']):
                    st.image(row['截圖檔案'], caption=f"{row['名稱']} ({row['狀態']})", use_container_width=True)
                else:
                    st.warning(f"無法取得 {row['名稱']} 的截圖")

        # Download buttons
        col1, col2 = st.columns(2)
        
        with col1:
            csv = df.drop(columns=['截圖檔案']).to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "下載檢測報告 (CSV)",
                csv,
                "website_status_report.csv",
                "text/csv",
                key='download-csv',
                use_container_width=True
            )
            
        with col2:
            docx_data = create_docx(df)
            st.download_button(
                label="下載 Word 完整報告 (含截圖)",
                data=docx_data,
                file_name=f"website_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key='download-docx',
                use_container_width=True
            )
